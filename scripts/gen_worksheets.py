#!/usr/bin/env python3

# Generate .docx/.xlsx/.pdf versions of candidate lists
# Paul "Worthless" Nijjar, 2025-12-30

import csv
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
import docx.shared
import os, sys
from datetime import datetime
import xlsxwriter

# For formatting docx cells
CANDIDATE_NAME_INCHES = 2.0
CANDIDATE_NOTE_INCHES = 5.0

DATADIR="docs/_data"
TEST_POSITION='SchoolBoard-Public-English-Kitchener'
TEST_WARD='Kitchener-Ward-09'
OUTDIR="docs/worksheets"

# I am not happy with reading and appending into a giant list, but
# whatever. The numbers are small enough that inefficiencies should
# not matter. 

# Pandas would probably be better but I can't use NumPy on this
# laptop without yak-shaving. 

static_text = {
  'intro': "Use this handy sheet to figure out the candidates"
      " you want to vote for. ",
  'datestamp': "Worksheet generated on {}".format(
        datetime.now().strftime("%A, %B %e, %Y, %l:%M%P")
        ),
  'title': "Municipal Election Candidate Sheet",
  'candidate_name': "Candidate Name",
  'candidate_notes': "My Notes",
}

# ---- HELPER FUNCTIONS

# From: https://stackoverflow.com/questions/37757203/making-cells-bold-in-a-table-using-python-docx
def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


# Doesn't work?
def make_rows_vertically_centred(*rows):
    for row in rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# https://stackoverflow.com/questions/43051462/python-docx-how-to-set-cell-width-in-tables
def set_column_width_in_inches(column, length_in_inches):
    column.width = docx.shared.Inches(length_in_inches)

    for cell in column.cells:
        cell.width = docx.shared.Inches(length_in_inches)


# https://stackoverflow.com/questions/41578438/how-do-you-keep-table-rows-together-in-python-docx
# SO UGLY
def keep_tables_on_one_page(doc):
    tags = doc.element.xpath('//w:tr[position() < last()]/w:tc/w:p')
    for tag in tags:
        ppr = tag.get_or_add_pPr()
        ppr.keepNext_val = True


# Write a line to the given file descriptor, with given number of
# newlines
def writeln(fd, str, num_newlines = 1):
    fd.write("{}{}".format(str, (num_newlines * "\n")))


# Get candidate name including withdrawn status. Does not bold
# anything.
def get_candidate_name(nom):
    candidate_text = "{} {}".format(
        nom['Given_Names'],
        nom['Last_Name'],
        )
    if nom['Withdrawn']:
        candidate_text = "{} (WITHDRAWN)".format(candidate_text)

    return candidate_text


# ---- GENERATE DOCUMENT FUNCTIONS 

# ward: unique identifier (eg 'Kitchener-Ward-09')
# pos_data: big complicated structure
# races: list of positions, including _SELF

def gen_docx(ward, pos_data, races):
    d = Document()

    d.add_heading(static_text['title'], 0)
    d.add_paragraph(static_text['datestamp'])
    d.add_paragraph(static_text['intro'])

    race_list = races.split(",")

    for r in race_list:
        if r == '_SELF':
            r = ward

        head = d.add_heading("{}".format( pos_data[r]['desc']))
        head.paragraph_format.keep_with_next = True

        elected_text = "{} to be elected".format(
            pos_data[r]['num_to_elect']
            )
        if pos_data[r]['acclaimed']:
            elected_text = "{} : ACCLAIMED".format(elected_text)

        p = d.add_paragraph(elected_text)
        p.paragraph_format.keep_with_next = True

        t = d.add_table(rows=1, cols=2)
        header_cells = t.rows[0].cells
        header_cells[0].text = "{}".format(static_text['candidate_name'])
        header_cells[1].text = "{}".format(static_text['candidate_notes'])
        make_rows_vertically_centred(t.rows[0])
        make_rows_bold(t.rows[0])

        for nom in pos_data[r]['candidates']:
            row = t.add_row()
            row_cells = row.cells

            row_cells[0].text = get_candidate_name(nom)
            row_cells[1].text = ""
            make_rows_vertically_centred(row)

        set_column_width_in_inches(t.columns[0], CANDIDATE_NAME_INCHES)
        set_column_width_in_inches(t.columns[1], CANDIDATE_NOTE_INCHES)

    keep_tables_on_one_page(d)

    d.save(os.path.join(OUTDIR,"docx","{}.docx".format(ward)))




# ward: unique identifier (eg 'Kitchener-Ward-09')
# pos_data: big complicated structure
# races: list of positions, including _SELF

def gen_xlsx(ward, pos_data, races):

    workbook = xlsxwriter.Workbook(os.path.join(
        OUTDIR,
        "xlsx",
        "{}.xlsx".format(ward)
        ))
    # Add a bold format to use to highlight cells.
    bold = workbook.add_format({'bold': True})
    h1 = workbook.add_format({'font_size': 18, 'bold': True})
    h2 = workbook.add_format({'font_size': 15, 'bold': True})
    table_header = workbook.add_format({
        'bold': True,
        'border': 1,  # solid
        })

    # This will be bad when I want to use strikeout. 
    table_cell = workbook.add_format({
        'left': 1,  # solid
        'right': 1,  # solid
        'text_wrap': True,
        'valign': 'vcenter',
        })

    table_bottom = workbook.add_format({
        'bottom': 1,
        'left': 1,  # solid
        'right': 1,  # solid
        'text_wrap': True,
        'valign': 'vcenter',
        })


    spreadsheet = workbook.add_worksheet('Candidates')
    row = 0

    spreadsheet.set_column(0, 0, 30)
    spreadsheet.set_column(1, 1, 50)
    
    spreadsheet.write(row, 0, static_text['title'], h1)
    row += 2

    spreadsheet.write(row, 0, static_text['datestamp'])
    row += 1
    spreadsheet.write(row, 0, static_text['intro'])
    row += 1 

    race_list = races.split(",")

    for r in race_list:
        row += 2

        if r == '_SELF':
            r = ward

        spreadsheet.write(row, 0, "{}".format( pos_data[r]['desc']), h2)
        row += 1

        elected_text = "{} to be elected".format(
            pos_data[r]['num_to_elect']
            )
        if pos_data[r]['acclaimed']:
            elected_text = "{} : ACCLAIMED".format(elected_text)

        spreadsheet.write(row, 0, elected_text)
        row += 2

        spreadsheet.write(
            row, 
            0, 
            "{}".format(static_text['candidate_name']),
            table_header,
            )
        spreadsheet.write(
            row, 
            1, 
            "{}".format(static_text['candidate_notes']),
            table_header,
            )
        row += 1

        candidates_so_far = 0
        num_candidates = pos_data[r]['num_candidates']

        for nom in pos_data[r]['candidates']:

            # We cannot set the format of a cell later without 
            # knowing its contents. 
            if candidates_so_far + 1 == num_candidates:
                format = table_bottom
            else:
                format = table_cell

            spreadsheet.write(
                row, 
                0, 
                get_candidate_name(nom),
                format,
                )
            spreadsheet.write_blank(
                row,
                1,
                "",
                format,
                )
            
            row += 1 
            candidates_so_far += 1

    workbook.close()


# ward: unique identifier (eg 'Kitchener-Ward-09')
# pos_data: big complicated structure
# races: list of positions, including _SELF

def gen_plaintext(ward, pos_data, races):

    with open (
        os.path.join(OUTDIR,"text","{}.txt".format(ward)),
        'w',
        encoding="utf-8",
        ) as d: 

        writeln(d, static_text['title'])
        writeln(d, ("=" * len(static_text['title'])), 2)


        writeln(d, static_text['datestamp'], 2)
        writeln(d, static_text['intro'], 2)

        race_list = races.split(",")

        for r in race_list:
            if r == '_SELF':
                r = ward

            writeln(d, pos_data[r]['desc'])
            writeln(d, ("-" * len(pos_data[r]['desc'])), 2)


            elected_text = "{} to be elected".format(
                pos_data[r]['num_to_elect']
                )
            if pos_data[r]['acclaimed']:
                elected_text = "{} : ACCLAIMED".format(elected_text)

            writeln(d, elected_text, 2)

            for nom in pos_data[r]['candidates']:
                writeln(d, "- {}".format(get_candidate_name(nom)))

            writeln(d, "", 2)

    d.close()



# --- READ DATA 



# Ugh I hate nested structures
pos_data = {}

"""
  'desc' => '',
  'candidates' => [],
  'num_to_elect' => 1,
  'ward' => '',
  'acclaimed' => False,
  'num_candidates' => 1,
"""

print("Current directory is: {}".format(os.getcwd()))
 


with open (os.path.join(DATADIR,"internal/position-tags.csv")) as r:
    pos_csv = csv.DictReader(r)

    for row in pos_csv:
        pos_name = row['PositionUniqueName']
        pos_data[pos_name] = {}
        pos_data[pos_name]['desc'] = row['PositionDesc']
        pos_data[pos_name]['candidates'] = []
        pos_data[pos_name]['num_to_elect'] = int(row['NumberToElect'])
        pos_data[pos_name]['ward_municipality'] = row['WardMunicipality']


with open(os.path.join(DATADIR,"sync/nominees.csv")) as r:
    nom_csv = csv.DictReader(r)

    for row in nom_csv:
        pos_data[row['PositionUniqueName']]['candidates'].append(row)

municipality_map = []
with open (os.path.join(DATADIR,"internal/municipality-map.csv")) as r:
    mun_csv = csv.DictReader(r)

    for row in mun_csv:
        municipality_map.append(row)


# Need to sort nominees by last name? 
for pos in pos_data:
  pos_data[pos]['candidates'].sort(
      key = lambda x: "{},{}".format(x['Last_Name'],x['Given_Names'])
      )

  pos_data[pos]['num_candidates'] = len(pos_data[pos]['candidates'])
  pos_data[pos]['acclaimed'] = (
      pos_data[pos]['num_candidates'] <= pos_data[pos]['num_to_elect']
      )
      



# ---- GEN DOCS


for ward in pos_data:
    if pos_data[ward]['ward_municipality'] == 'N/A':
        continue

    races = "NOT-FOUND"

    for m in municipality_map:
        if m['Name'] == pos_data[ward]['ward_municipality']:
            races = m['Races']
            break
      
    if races == 'NOT-FOUND':
        raise NameError

    gen_plaintext(ward, pos_data, races)
    gen_docx(ward, pos_data, races)
    gen_xlsx(ward, pos_data, races)





# ---- BAD TESTING

"""
print(pos_data[TEST_POSITION])

for nom in pos_data[TEST_POSITION]['candidates']:
    print("{} {}".format(nom['Given_Names'],nom['Last_Name']))
"""


